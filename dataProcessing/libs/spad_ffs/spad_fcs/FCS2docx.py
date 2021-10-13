from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from datetime import date
import ntpath
from binData import binData
from arrivalTimes2TimeTrace import arrivalTimes2Data
from addSumColumn import addSumColumn
from meas_to_count import file_to_count
from FCS2Corr import FCSLoadAndCorrSplit
from listFiles import listFiles
import matplotlib.pyplot as plt
import numpy as np
from plotAiry import plotAiry
from FCS2Corr import FCS2CorrSplit
from FCS2Corr import FCS2Corr
from FCS2Corr import plotFCScorrelations
from pathlib import Path
from os import  getcwd
from getFCSinfo import getFCSinfo
from corr2csv import corr2csv


def FCS2docx(folderName=[], Mtype='25ch'):
    
    """
    Analyze FCS data and store results in image files and a .docx file
    ===========================================================================
    Input           Meaning
    ---------------------------------------------------------------------------
    folderName      Folder name with FCS data
                    Leave blank for current folder
    ===========================================================================

    ===========================================================================
    Output
    ---------------------------------------------------------------------------
    images of time traces, autocorrelations and an overview .docx file
    ===========================================================================
    """
    
    # PARSE INPUT
    if folderName == []:
        folderName = getcwd()
    folderName = folderName.replace("\\", "/")
    folderName = Path(folderName)
    
    # OPEN WORD FILE
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.bottom_margin = Mm(25)
    section.top_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    
    # TITLE
    title = date.today().strftime("%Y-%m-%d")
    title = title + ' FCS analysis'
    document.add_heading(title, 0)
    
    # CHECK PICKLE FILES
    outp = listFiles(folderName, 'bin')
    
    # GO THROUGH EACH FILE
    for file in outp:
        if 'arrival_times' not in file:
            
            # file with photon arrival times found
            fileName = ntpath.basename(file)
            print('==========================================================')
            print('File found: ' + fileName)
            print('==========================================================')
            # ================================================================
            document.add_heading('File ' + fileName, level=1)
            # ================================================================
            
            # get pixel dwell time from text file (TO DO)
            FCSinfo = getFCSinfo(file[0:-4] + '_info.txt')
            #pxdwell = 1e-3 / ReadSpeed  # s
            pxdwell = FCSinfo.dwellTime
            
#            # open file
#            # data = arrivalTimes2Data(file)
            data = file_to_count(file)
            data = data[0]
            data = addSumColumn(data)
#            
#            # bin files for plotting time trace
            print('Binning data for plotting time trace.')
            binSize = 2000000
            binDuration = pxdwell * binSize
            Nrows = 1000000
#            # -----
            dataB = binData(data[0:Nrows,:], binSize)
            plt.figure()
            plt.plot(np.arange(0, binDuration*len(dataB), binDuration), dataB[:, -1])
            plt.xlabel('Time [s]')
            plt.ylabel('Total # of photons per ' + str(1000 * binDuration) + ' ms')
            plt.xlim([0, binDuration*len(dataB)])
            plt.title(fileName[0:30] + '...')
            plt.tight_layout()
            picName = fileName[0:-4] + '_timetrace.png'
#            plt.savefig(picName)
#            # ================================================================
#            document.add_heading('Time trace', level=2)
#            document.add_picture(picName, width=Inches(4))
#            # ================================================================
#            
#            # plot Airy pattern
#            print('Calculating Airy pattern.')
#            airy = plotAiry(data)
#            picName = fileName[0:-4] + '_airy.png'
#            plt.savefig(picName)
#            # ================================================================
#            document.add_heading('Airy pattern', level=2)
#            document.add_paragraph('Number of photons per channel over the entire measurement')
#            document.add_picture(picName, width=Inches(4))
#            # ================================================================
#            
#            # calculate autocorrelations
#            # if single element is used, only calculate autocorr of this element
#            # else do complete calculation
#            print('Calculating correlations.')
            # =================================================================
            document.add_heading('Correlations', level=2)
            # =================================================================
            zoom = int(1e-5 / pxdwell)
#            if np.max(airy[0:25]) / airy[25] > 0.99:
#                # single element used
#                det = int(np.argmax(airy[0:25]))
#                G = FCS2Corr(data, 1e6*pxdwell, [det])
#                Gsingle = getattr(G, 'det'+ str(det))
#                plotFCScorrelations(G, plotList='all')
#                plt.xlim(left=2*pxdwell)
#                plt.ylim(top=np.max(Gsingle[2:,1]))
#                picName = fileName[0:-4] + '_G.png'
#                plt.savefig(picName)
#                # =============================================================
#                document.add_heading('Autocorrelation single detector', level=3)
#                document.add_picture(picName, width=Inches(4))
#                # =============================================================
#                plt.xlim(left=zoom*pxdwell)
#                plt.ylim(top=np.max(Gsingle[zoom:,1]))
#                picName = fileName[0:-4] + '_Gzoom.png'
#                plt.savefig(picName)
#                # =============================================================
#                document.add_picture(picName, width=Inches(4))
                # =============================================================
            if Mtype == '2MPD':
                G = FCS2CorrSplit(data, 1e6*pxdwell, ['2MPD'], 50, 6)
                corr2csv(G, fileName[0:-4])
                plotFCScorrelations(G, ['auto1_average', 'auto2_average', 'cross12_average', 'cross21_average', 'cross_average'])
                plt.xlim(left=2*pxdwell)
                picName = fileName[0:-4] + '_crosscorr.png'
                plt.savefig(picName)
                # =============================================================
                document.add_heading('Correlations', level=3)
                document.add_picture(picName, width=Inches(4))
                # =============================================================
            else:
                plotList = ['central', 'sum3', 'sum5', 'chessboard', 'ullr']
                maxG1 = np.zeros([len(plotList), 1])
                maxG2 = np.zeros([len(plotList), 1])
                # G = FCS2CorrSplit(data, 1e6*pxdwell, plotList, 50, 5)
                G = FCSLoadAndCorrSplit(file, plotList, 16, 10)
                corr2csv(G, fileName[0:-4])
                plotList = ['central_average', 'sum3_average', 'sum5_average', 'chessboard_average', 'ullr_average']
                for i in range(len(plotList)):
                    Gsingle = getattr(G, plotList[i])
                    maxG1[i] = np.max(Gsingle[2:,1])
                    maxG2[i] = np.max(Gsingle[zoom:,1])
                
                plotFCScorrelations(G, plotList=['central_average', 'sum3_average', 'sum5_average'])
                plt.xlim(left=2*pxdwell)
                plt.ylim(top=np.max(maxG1[0:2]))
                picName = fileName[0:-4] + '_G135.png'
                plt.savefig(picName)
                # =============================================================
                document.add_heading('Autocorrelations', level=3)
                document.add_picture(picName, width=Inches(4))
                # =============================================================
                
                plt.xlim(left=zoom*pxdwell)
                plt.ylim(top=np.max(maxG2[0:2]))
                picName = fileName[0:-4] + '_G135_zoom.png'
                plt.savefig(picName)
                # =============================================================
                document.add_picture(picName, width=Inches(4))
                # =============================================================
                
                plotFCScorrelations(G, plotList=['sum5_average'])
                plt.xlim(left=zoom*pxdwell)
                plt.ylim(top=np.max(maxG2[2]))
                picName = fileName[0:-4] + '_sum5_zoom.png'
                plt.savefig(picName)
                # =============================================================
                document.add_picture(picName, width=Inches(4))
                # =============================================================
                
                #plotFCScorrelations(G, plotList=['sum5', 'det-15'])
                #plt.xlim(left=zoom*pxdwell)
                #plt.ylim(top=np.max(maxG2[2:4]))
                #picName = fileName[0:-4] + '_sum5MinusHotPixel_zoom.png'
                #plt.savefig(picName)
                # =============================================================
                #document.add_paragraph('det-15 means the sum of all detector elements except for pixel 15 (hot pixel).')
                #document.add_picture(picName, width=Inches(4))
                # =============================================================
                
                
                plotFCScorrelations(G, plotList=['chessboard_average', 'ullr_average'])
                plt.xlim(left=pxdwell)
                plt.ylim(top=np.max(maxG2[4:]))
                picName = fileName[0:-4] + '_crosscorr_zoom.png'
                plt.savefig(picName)
                # =============================================================
                document.add_heading('Cross correlations', level=3)
                document.add_paragraph('Chessboard means the cross-correlation between the even number pixels and the odd numbered pixels. ULLR means the cross-correlation between the top-left and the bottom-right triangle (UpperLeft-LowerRight).')
                document.add_picture(picName, width=Inches(4))
                # =============================================================
            
    
    document.save('Overview_results.docx')